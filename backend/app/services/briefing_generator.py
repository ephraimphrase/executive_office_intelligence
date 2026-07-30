import io
import logging
from datetime import date as date_cls
from datetime import datetime

from agents.briefing_agent import BriefingAgent
from app.integrations.openai_client import get_openai_client

logger = logging.getLogger(__name__)

class BriefingPack:
    def __init__(self, **kwargs):
        for k, v in kwargs.items():
            setattr(self, k, v)

class BriefingGeneratorService:
    def __init__(self):
        self.agent = BriefingAgent(get_openai_client())

    async def generate_briefing(self, date: str, db) -> BriefingPack:
        """Assemble the full Executive Briefing Pack: today's meetings (with
        participants), outstanding tasks, critical emails, open risks,
        pending decisions, travel plans, weather, and AI talking points —
        everything the spec's Executive Briefing Generator calls for."""
        logger.info(f"Generating briefing for {date}")
        events, tasks, emails, risks, decisions, travel = [], [], [], [], [], []
        weather: dict = {}

        target_date = date if isinstance(date, date_cls) else datetime.strptime(date, "%Y-%m-%d").date()

        if db is not None:
            from sqlalchemy import select

            from app.models.decision import Decision, DecisionStatus
            from app.models.email_record import EmailRecord, EmailStatus
            from app.models.event import Event
            from app.models.risk import Risk, RiskSeverity, RiskStatus
            from app.models.task import Task, TaskStatus

            day_start = datetime.combine(target_date, datetime.min.time())
            day_end = datetime.combine(target_date, datetime.max.time())

            events_res = await db.execute(
                select(Event)
                .where(Event.start_datetime >= day_start, Event.start_datetime <= day_end)
                .order_by(Event.start_datetime)
            )
            all_events = events_res.scalars().all()
            events = [
                {
                    "title": e.title,
                    "start_datetime": e.start_datetime,
                    "end_datetime": e.end_datetime,
                    "location": e.location,
                    "event_type": e.event_type.value if hasattr(e.event_type, "value") else str(e.event_type),
                    "attendees": e.attendees or [],
                    "agenda": e.agenda,
                    "board_paper_required": e.board_paper_required,
                    "preparation_required": e.preparation_required,
                }
                for e in all_events
            ]
            travel = [e for e in events if e["event_type"] in ("TRAVEL", "SITE_VISIT")]

            tasks_res = await db.execute(
                select(Task).where(Task.status != TaskStatus.DONE).order_by(Task.due_date).limit(10)
            )
            tasks = [
                {
                    "title": t.title,
                    "due_date": t.due_date,
                    "priority": t.priority.value if hasattr(t.priority, "value") else str(t.priority),
                }
                for t in tasks_res.scalars().all()
            ]

            email_res = await db.execute(
                select(EmailRecord)
                .where(EmailRecord.is_high_priority == True, EmailRecord.status != EmailStatus.ARCHIVED)  # noqa: E712
                .order_by(EmailRecord.received_at.desc())
                .limit(10)
            )
            emails = [
                {
                    "subject": e.subject,
                    "sender": e.sender_name or e.sender_email,
                    "summary": e.ai_summary,
                    "received_at": e.received_at,
                    "requires_response": e.status not in (EmailStatus.REPLIED,),
                }
                for e in email_res.scalars().all()
            ]

            risk_res = await db.execute(
                select(Risk).where(
                    Risk.status == RiskStatus.OPEN,
                    Risk.severity.in_([RiskSeverity.CRITICAL, RiskSeverity.HIGH]),
                )
            )
            risks = [
                {"description": r.description, "severity": r.severity.value, "category": r.category}
                for r in risk_res.scalars().all()
            ]

            decision_res = await db.execute(
                select(Decision).where(Decision.status == DecisionStatus.PENDING_IMPLEMENTATION).limit(10)
            )
            decisions = [
                {"description": d.description, "made_by": d.made_by, "deadline": d.deadline}
                for d in decision_res.scalars().all()
            ]

        try:
            from app.config import get_settings
            from app.integrations.weather import get_current_weather

            settings = get_settings()
            weather = await get_current_weather(settings.gvp_location_lat, settings.gvp_location_lon)
        except Exception as e:
            logger.warning(f"Weather lookup failed for briefing: {e}")
            weather = {}

        talking_points: list = []
        if events:
            talking_points = await self.agent.generate_talking_points(
                {"title": "Today's key meetings", "events": events[:5]},
                {"tasks": tasks, "risks": risks},
            )

        risk_summary = await self.agent.generate_risk_summary(risks) if risks else ""
        priorities = await self.agent.generate_daily_priorities(events, tasks, emails, risks)
        narrative = await self.agent.generate_executive_narrative({"date": str(date), "priorities": priorities})

        return BriefingPack(
            date=date, narrative=narrative, priorities=priorities, events=events, tasks=tasks,
            emails=emails, risks=risks, risk_summary=risk_summary, decisions=decisions,
            travel=travel, weather=weather, talking_points=talking_points,
        )

    async def get_today_briefing(self, db) -> BriefingPack:
        return await self.generate_briefing(date_cls.today(), db)

    def _extract_events(self, briefing) -> list:
        events = getattr(briefing, "events", None)
        if events is not None:
            return events
        events_summary = getattr(briefing, "events_summary", None) or {}
        return events_summary.get("events", []) if isinstance(events_summary, dict) else []

    def _extract_narrative(self, briefing) -> str:
        narrative = getattr(briefing, "narrative", None)
        if narrative:
            return narrative
        full_content = getattr(briefing, "full_content", None) or {}
        return full_content.get("narrative", "") if isinstance(full_content, dict) else ""

    async def export_to_docx(self, briefing, db) -> bytes:
        from docx import Document as DocxDocument

        from app.services.word_generator import WordGeneratorService

        events = self._extract_events(briefing)
        narrative = self._extract_narrative(briefing)
        date_str = str(getattr(briefing, "date", ""))

        doc = DocxDocument()
        doc.add_heading(f"Executive Briefing — {date_str}", level=1)
        if narrative:
            doc.add_paragraph(narrative)
        doc.add_heading("Today's Schedule", level=2)

        word_svc = WordGeneratorService()
        word_svc._add_schedule_table(doc, events)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def export_to_pptx(self, briefing, db) -> bytes:
        from app.services.pptx_generator import PowerPointGeneratorService

        events = self._extract_events(briefing)
        pack_dict = {
            "date": getattr(briefing, "date", ""),
            "priorities": getattr(briefing, "suggested_priorities", None) or getattr(briefing, "priorities", []) or [],
            "events": events,
            "risks": (getattr(briefing, "risks_summary", None) or {}).get("risks", []) if isinstance(getattr(briefing, "risks_summary", None), dict) else getattr(briefing, "risks", []) or [],
            "decisions": (getattr(briefing, "pending_decisions", None) or {}).get("decisions", []) if isinstance(getattr(briefing, "pending_decisions", None), dict) else getattr(briefing, "decisions", []) or [],
            "talking_points": getattr(briefing, "talking_points", []) or [],
        }

        pptx_svc = PowerPointGeneratorService()
        return await pptx_svc.generate_briefing_deck(pack_dict, db)

    async def export_to_pdf(self, briefing, db) -> bytes:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        events = self._extract_events(briefing)
        narrative = self._extract_narrative(briefing)
        date_str = str(getattr(briefing, "date", ""))

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph(f"Executive Briefing — {date_str}", styles["Title"]), Spacer(1, 12)]
        if narrative:
            story.append(Paragraph(narrative, styles["Normal"]))
            story.append(Spacer(1, 12))

        story.append(Paragraph("Today's Schedule", styles["Heading2"]))
        table_data = [["Time", "Title", "Location"]]
        for e in events:
            start = e.get("start_datetime") or e.get("start_time") or ""
            title = e.get("title") or e.get("agenda") or ""
            location = e.get("location") or e.get("venue") or ""
            table_data.append([str(start), str(title), str(location)])

        table = Table(table_data)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E3B4E")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(table)

        doc.build(story)
        return buf.getvalue()

    async def send_briefing_email(self, briefing, recipients: list) -> bool:
        from app.integrations.microsoft_graph import MicrosoftGraphClient

        logger.info(f"Sending briefing to {recipients}")
        graph = MicrosoftGraphClient()
        narrative = self._extract_narrative(briefing)
        date_str = str(getattr(briefing, "date", ""))
        body = narrative or f"Executive briefing for {date_str}"

        ok = True
        for recipient in recipients:
            sent = await graph.send_email(recipient, f"Executive Briefing — {date_str}", body)
            ok = ok and sent
        return ok
