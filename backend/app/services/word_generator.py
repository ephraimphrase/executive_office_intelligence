import io
import logging

logger = logging.getLogger(__name__)


def _fmt_time(value) -> str:
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%H:%M")
    return str(value)


def _fmt_date(value) -> str:
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y-%m-%d")
    return str(value)


class WordGeneratorService:
    """Generates real .docx files matching the GVP's Executive Daily Schedule template
    (Date, Time, Agenda, Venue, Notes, Owner columns)."""

    def _new_document(self, title: str):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document()
        heading = doc.add_heading("Group Vice President — Executive Daily Schedule", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph(title)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].italic = True
        subtitle.runs[0].font.size = Pt(11)
        return doc

    def _add_schedule_table(self, doc, events: list):
        headers = ["Date", "Time", "Agenda", "Venue", "Notes", "Owner"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for event in events:
            row = table.add_row().cells
            row[0].text = _fmt_date(event.get("date") or event.get("start_datetime"))
            start = _fmt_time(event.get("start_time") or event.get("start_datetime"))
            end = _fmt_time(event.get("end_time") or event.get("end_datetime"))
            row[1].text = f"{start} - {end}" if start or end else ""
            row[2].text = event.get("agenda") or event.get("title") or ""
            row[3].text = event.get("venue") or event.get("location") or ""
            row[4].text = event.get("notes") or ""
            row[5].text = event.get("owner") or event.get("owner_name") or ""

        return table

    def _to_bytes(self, doc) -> bytes:
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def generate_daily_schedule(self, date: str, events: list, db) -> bytes:
        logger.info(f"Generating daily schedule DOCX for {date}")
        doc = self._new_document(f"Daily Schedule — {date}")
        self._add_schedule_table(doc, events)
        return self._to_bytes(doc)

    async def generate_weekly_schedule(self, week_start: str, db, events: list | None = None) -> bytes:
        doc = self._new_document(f"Weekly Schedule — Week of {week_start}")
        self._add_schedule_table(doc, events or [])
        return self._to_bytes(doc)

    async def generate_meeting_minutes(self, meeting_record: dict, db) -> bytes:
        from docx.shared import Pt

        doc = self._new_document(f"Meeting Minutes — {meeting_record.get('title', 'Meeting')}")
        doc.add_paragraph(f"Date: {_fmt_date(meeting_record.get('meeting_date'))}")
        doc.add_paragraph(f"Chairperson: {meeting_record.get('chairperson', '')}")

        doc.add_heading("Attendees", level=2)
        for attendee in meeting_record.get("participants", []) or []:
            doc.add_paragraph(str(attendee), style="List Bullet")

        doc.add_heading("Minutes", level=2)
        doc.add_paragraph(meeting_record.get("minutes") or meeting_record.get("ai_minutes") or "")

        doc.add_heading("Action Items", level=2)
        for action in meeting_record.get("action_items", []) or []:
            text = action.get("description") if isinstance(action, dict) else str(action)
            doc.add_paragraph(text, style="List Bullet")

        return self._to_bytes(doc)

    async def generate_action_register(self, tasks: list, db) -> bytes:
        doc = self._new_document("Action Register")
        headers = ["Title", "Owner", "Due Date", "Priority", "Status"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for task in tasks:
            row = table.add_row().cells
            row[0].text = task.get("title", "")
            row[1].text = task.get("owner", "") or task.get("assigned_to", "")
            row[2].text = _fmt_date(task.get("due_date"))
            row[3].text = task.get("priority", "")
            row[4].text = task.get("status", "")

        return self._to_bytes(doc)

    async def generate_decision_register(self, decisions: list, db) -> bytes:
        doc = self._new_document("Executive Decision Register")
        headers = ["Decision", "Made By", "Date", "Status", "Responsible Person"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for decision in decisions:
            row = table.add_row().cells
            row[0].text = decision.get("description", "")
            row[1].text = decision.get("made_by", "")
            row[2].text = _fmt_date(decision.get("decision_date"))
            row[3].text = decision.get("status", "")
            row[4].text = decision.get("responsible_person", "")

        return self._to_bytes(doc)
